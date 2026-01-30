"""
DOCX Exporter for Stage E (Export).

Exports pipeline results to Microsoft Word format using python-docx.
Supports AlignmentLayer export with structured document generation.

Module Version: 2.0.0
"""

import logging
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Inches = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None
    WD_STYLE_TYPE = None

from pydantic import BaseModel

from ...schemas.export import (
    ExportFormat,
    ExportOptions,
    ExportSpec,
)
from ...schemas.alignment_layer import (
    AlignmentLayer,
    LaTeXAlignment,
    MatchedPairV2,
    VerificationStatus,
)
from .base import BaseExporter, ExporterConfig
from ..exceptions import ExporterError
from ..builders import AlignmentDocumentBuilder, DocumentSection, DocumentStructure

logger = logging.getLogger(__name__)


# =============================================================================
# Configuration
# =============================================================================

@dataclass
class DOCXExporterConfig(ExporterConfig):
    """Configuration for DOCX exporter.

    Attributes:
        page_width_inches: Page width in inches
        page_height_inches: Page height in inches
        margin_inches: Page margin in inches (uniform)
        font_name: Default font name
        font_size_pt: Default font size in points
        heading_font_name: Font for headings
        include_toc: Include table of contents
        include_header: Include document header
        include_footer: Include page numbers in footer
    """
    page_width_inches: float = 8.5
    page_height_inches: float = 11.0
    margin_inches: float = 1.0
    font_name: str = "Times New Roman"
    font_size_pt: int = 11
    heading_font_name: str = "Arial"
    include_toc: bool = False
    include_header: bool = True
    include_footer: bool = True


# =============================================================================
# DOCX Exporter
# =============================================================================

class DOCXExporter(BaseExporter[DOCXExporterConfig]):
    """Export pipeline results to DOCX format.

    Stage E exporter that uses python-docx library to generate
    Microsoft Word documents. Primary input: AlignmentLayer.

    Supports:
    - AlignmentLayer export (primary use case)
    - LaTeX-Visual alignment sections
    - HITL review sections
    - Quality metrics tables
    - Legacy data types (dict, list, BaseModel)

    Usage:
        exporter = DOCXExporter()

        # AlignmentLayer export (Stage E primary)
        spec = exporter.export(alignment_layer, options, alignment_layer.id)

        # Or export to bytes
        docx_bytes = exporter.export_to_bytes(alignment_layer, options)

    Requires:
        python-docx>=0.8.11
    """

    def __init__(self, config: Optional[DOCXExporterConfig] = None):
        """Initialize DOCX exporter.

        Args:
            config: Exporter configuration. Uses defaults if None.

        Raises:
            ImportError: If python-docx is not installed.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )
        super().__init__(config)

    def _default_config(self) -> DOCXExporterConfig:
        """Create default configuration."""
        return DOCXExporterConfig()

    @property
    def format(self) -> ExportFormat:
        """Get export format."""
        return ExportFormat.DOCX

    @property
    def content_type(self) -> str:
        """Get MIME content type."""
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def file_extension(self) -> str:
        """Get file extension."""
        return ".docx"

    def _create_document(self) -> Any:
        """Create a new Word document with default styling.

        Returns:
            Document object
        """
        doc = Document()

        # Configure default paragraph style
        style = doc.styles['Normal']
        font = style.font
        font.name = self.config.font_name
        font.size = Pt(self.config.font_size_pt)

        return doc

    def _add_title(self, doc: Any, title: str) -> None:
        """Add document title.

        Args:
            doc: Document object
            title: Title text
        """
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_metadata_section(
        self,
        doc: Any,
        data: Any,
        options: ExportOptions,
    ) -> None:
        """Add metadata section to document.

        Args:
            doc: Document object
            data: Source data
            options: Export options
        """
        if not options.include_metadata:
            return

        doc.add_heading("Metadata", level=1)

        # Add timestamp
        doc.add_paragraph(
            f"Generated: {datetime.now(timezone.utc).isoformat()}"
        )

        # Add data attributes
        if hasattr(data, "__dict__"):
            for key, value in data.__dict__.items():
                if key.startswith("_"):
                    continue
                if isinstance(value, (str, int, float, bool)):
                    doc.add_paragraph(f"{key}: {value}")
        elif isinstance(data, BaseModel):
            for key, value in data.model_dump().items():
                if key.startswith("_"):
                    continue
                if isinstance(value, (str, int, float, bool)):
                    doc.add_paragraph(f"{key}: {value}")

    def _add_content_section(
        self,
        doc: Any,
        data: Any,
        options: ExportOptions,
    ) -> int:
        """Add main content section to document.

        Args:
            doc: Document object
            data: Source data
            options: Export options

        Returns:
            Number of elements added
        """
        element_count = 0

        # Handle different data types
        if hasattr(data, "outputs"):
            doc.add_heading("Outputs", level=1)
            for i, output in enumerate(data.outputs):
                element_count += 1
                self._add_output_item(doc, output, i + 1)

        elif hasattr(data, "nodes"):
            doc.add_heading("Graph Nodes", level=1)
            for i, node in enumerate(data.nodes):
                element_count += 1
                self._add_node_item(doc, node, i + 1)

        elif hasattr(data, "equations"):
            doc.add_heading("Equations", level=1)
            for i, equation in enumerate(data.equations):
                element_count += 1
                self._add_equation_item(doc, equation, i + 1)

        elif isinstance(data, dict):
            doc.add_heading("Data", level=1)
            for key, value in data.items():
                element_count += 1
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))

        elif isinstance(data, list):
            doc.add_heading("Items", level=1)
            for i, item in enumerate(data):
                element_count += 1
                doc.add_paragraph(f"{i + 1}. {item}")

        return element_count

    def _add_output_item(self, doc: Any, output: Any, index: int) -> None:
        """Add an output item to the document.

        Args:
            doc: Document object
            output: Output data
            index: Item index
        """
        # Add subheading
        doc.add_heading(f"Output {index}", level=2)

        # Add content based on output structure
        if hasattr(output, "content"):
            p = doc.add_paragraph()
            p.add_run("Content: ").bold = True
            p.add_run(str(output.content))

        if hasattr(output, "latex"):
            p = doc.add_paragraph()
            p.add_run("LaTeX: ").bold = True
            p.add_run(str(output.latex))

        if hasattr(output, "confidence"):
            p = doc.add_paragraph()
            p.add_run("Confidence: ").bold = True
            p.add_run(f"{output.confidence:.2%}")

    def _add_node_item(self, doc: Any, node: Any, index: int) -> None:
        """Add a graph node to the document.

        Args:
            doc: Document object
            node: Node data
            index: Item index
        """
        node_type = getattr(node, "type", "Unknown")
        doc.add_heading(f"Node {index}: {node_type}", level=2)

        if hasattr(node, "content"):
            doc.add_paragraph(str(node.content))

        if hasattr(node, "attributes"):
            for key, value in node.attributes.items():
                p = doc.add_paragraph()
                p.add_run(f"  {key}: ").italic = True
                p.add_run(str(value))

    def _add_equation_item(self, doc: Any, equation: Any, index: int) -> None:
        """Add an equation to the document.

        Args:
            doc: Document object
            equation: Equation data
            index: Item index
        """
        doc.add_heading(f"Equation {index}", level=2)

        if hasattr(equation, "latex"):
            p = doc.add_paragraph()
            p.add_run("LaTeX: ").bold = True
            # Use monospace-style formatting for LaTeX
            latex_run = p.add_run(str(equation.latex))
            latex_run.font.name = "Courier New"

        if hasattr(equation, "text"):
            p = doc.add_paragraph()
            p.add_run("Text: ").bold = True
            p.add_run(str(equation.text))

    def _add_provenance_section(
        self,
        doc: Any,
        data: Any,
        options: ExportOptions,
    ) -> None:
        """Add provenance chain section.

        Args:
            doc: Document object
            data: Source data
            options: Export options
        """
        if not options.include_provenance:
            return

        if not hasattr(data, "provenance"):
            return

        doc.add_heading("Provenance", level=1)
        provenance = data.provenance

        if hasattr(provenance, "created_at"):
            doc.add_paragraph(f"Created: {provenance.created_at}")

        if hasattr(provenance, "pipeline_version"):
            doc.add_paragraph(f"Pipeline Version: {provenance.pipeline_version}")

        if hasattr(provenance, "stages_completed"):
            doc.add_paragraph(f"Stages: {', '.join(provenance.stages_completed)}")

    # =========================================================================
    # AlignmentLayer Export Methods (Stage E)
    # =========================================================================

    def _is_alignment_layer(self, data: Any) -> bool:
        """Check if data is an AlignmentLayer instance."""
        return isinstance(data, AlignmentLayer)

    def _add_alignment_layer_content(
        self,
        doc: Any,
        layer: AlignmentLayer,
        options: ExportOptions,
    ) -> int:
        """Add AlignmentLayer content to document.

        Args:
            doc: Document object
            layer: AlignmentLayer to export
            options: Export options

        Returns:
            Number of elements added
        """
        builder = AlignmentDocumentBuilder(layer)
        doc_structure = builder.build()
        element_count = 0

        # Process each section from builder
        for section in doc_structure.sections:
            element_count += self._add_document_section(doc, section)

        return element_count

    def _add_document_section(self, doc: Any, section: DocumentSection) -> int:
        """Recursively add DocumentSection to DOCX.

        Args:
            doc: Document object
            section: DocumentSection to add

        Returns:
            Number of elements added
        """
        element_count = 1

        # Add section heading
        doc.add_heading(section.title, level=section.level)

        # Add content based on section type
        section_type = section.metadata.get("section_type", "")

        if section_type == "summary":
            self._add_summary_content(doc, section.content)
        elif section_type == "alignments":
            self._add_alignments_overview(doc, section.content)
        elif section_type == "matched_pairs":
            self._add_matched_pairs_content(doc, section.content)
        elif section_type == "hitl":
            self._add_hitl_overview(doc, section.content)
        elif section_type == "quality":
            self._add_quality_metrics_table(doc, section.content)
        elif section_type == "provenance":
            self._add_provenance_content(doc, section.content)
        elif section.content:
            # Generic content handling
            self._add_generic_content(doc, section.content)

        # Process subsections
        for subsection in section.subsections:
            element_count += self._add_document_section(doc, subsection)

        return element_count

    def _add_summary_content(self, doc: Any, content: Dict[str, Any]) -> None:
        """Add summary section content."""
        # Create summary table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        metrics = [
            ("Total Alignments", content.get("total_alignments", 0)),
            ("Verified", content.get("verified_count", 0)),
            ("Mismatches", content.get("mismatch_count", 0)),
            ("HITL Required", content.get("hitl_required_count", 0)),
            ("Avg. Alignment Score", f"{content.get('average_alignment_score', 0):.4f}"),
            ("Overall Quality", f"{content.get('overall_quality', 0):.4f}"),
            ("Needs HITL", "Yes" if content.get("needs_hitl") else "No"),
        ]

        for metric, value in metrics:
            row = table.add_row()
            row.cells[0].text = metric
            row.cells[1].text = str(value)

        # Verification breakdown
        breakdown = content.get("verification_breakdown", {})
        if breakdown:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Verification Breakdown: ").bold = True
            for status, count in breakdown.items():
                doc.add_paragraph(f"  - {status}: {count}")

    def _add_alignments_overview(self, doc: Any, content: Dict[str, Any]) -> None:
        """Add alignments overview."""
        count = content.get("count", 0)
        doc.add_paragraph(f"Total alignments: {count}")

    def _add_latex_alignment(self, doc: Any, alignment_data: Dict[str, Any]) -> None:
        """Add a single LaTeX alignment entry.

        Args:
            doc: Document object
            alignment_data: Alignment data dictionary
        """
        # LaTeX element
        latex = alignment_data.get("latex_element", {})
        p = doc.add_paragraph()
        p.add_run("LaTeX: ").bold = True
        latex_run = p.add_run(latex.get("latex_content", ""))
        latex_run.font.name = "Courier New"

        p = doc.add_paragraph()
        p.add_run("Type: ").bold = True
        p.add_run(latex.get("element_type", "unknown"))

        # Visual elements
        visuals = alignment_data.get("visual_elements", [])
        if visuals:
            p = doc.add_paragraph()
            p.add_run(f"Visual Elements ({len(visuals)}): ").bold = True
            for ve in visuals:
                doc.add_paragraph(
                    f"  - {ve.get('semantic_label', 'N/A')} "
                    f"[{ve.get('element_class', 'N/A')}] "
                    f"(conf: {ve.get('confidence', 0):.2%})"
                )

        # Scores
        p = doc.add_paragraph()
        p.add_run("Alignment Score: ").bold = True
        p.add_run(f"{alignment_data.get('alignment_score', 0):.4f}")

        p = doc.add_paragraph()
        p.add_run("Status: ").bold = True
        status = alignment_data.get("verification_status", "unknown")
        status_run = p.add_run(status.upper())
        if status == "verified":
            status_run.bold = True
        elif status == "mismatch":
            status_run.italic = True

    def _add_matched_pairs_content(self, doc: Any, content: Dict[str, Any]) -> None:
        """Add matched pairs content."""
        pairs = content.get("pairs", [])
        if not pairs:
            doc.add_paragraph("No matched pairs.")
            return

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        headers = ["LaTeX ID", "Visual ID", "Confidence", "Reason", "Verified"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for pair in pairs:
            row = table.add_row()
            row.cells[0].text = pair.get("latex_id", "")[:20]
            row.cells[1].text = pair.get("visual_id", "")[:20]
            row.cells[2].text = f"{pair.get('match_confidence', 0):.2%}"
            row.cells[3].text = pair.get("match_reason", "")[:30]
            row.cells[4].text = "Yes" if pair.get("verified") else "No"

    def _add_hitl_overview(self, doc: Any, content: Dict[str, Any]) -> None:
        """Add HITL overview content."""
        p = doc.add_paragraph()
        p.add_run("Flagged Items: ").bold = True
        p.add_run(str(content.get("flagged_count", 0)))

        p = doc.add_paragraph()
        p.add_run("Feedbacks Received: ").bold = True
        p.add_run(str(content.get("feedback_count", 0)))

        p = doc.add_paragraph()
        p.add_run("Pending Review: ").bold = True
        p.add_run(str(content.get("pending_count", 0)))

    def _add_quality_metrics_table(self, doc: Any, content: Dict[str, Any]) -> None:
        """Add quality metrics as a formatted table.

        Args:
            doc: Document object
            content: Quality metrics dictionary
        """
        # Main quality metrics table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Quality Metric"
        hdr_cells[1].text = "Value"
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True

        metrics = [
            ("Overall Quality", f"{content.get('overall_quality', 0):.4f}"),
            ("Avg. Alignment Score", f"{content.get('average_alignment_score', 0):.4f}"),
            ("Verified Ratio", f"{content.get('verified_ratio', 0):.2%}"),
            ("Mismatch Ratio", f"{content.get('mismatch_ratio', 0):.2%}"),
            ("Review Required", "Yes" if content.get("review_required") else "No"),
        ]

        if content.get("review_severity"):
            metrics.append(("Review Severity", content["review_severity"]))

        for metric, value in metrics:
            row = table.add_row()
            row.cells[0].text = metric
            row.cells[1].text = str(value)

        # Score distribution
        dist = content.get("score_distribution", {})
        if dist:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Score Distribution: ").bold = True
            doc.add_paragraph(f"  Min: {dist.get('min', 0):.4f}")
            doc.add_paragraph(f"  Max: {dist.get('max', 0):.4f}")
            doc.add_paragraph(f"  Avg: {dist.get('avg', 0):.4f}")

        # Calibration status
        calibration = content.get("calibration")
        if calibration:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Calibration Status: ").bold = True
            doc.add_paragraph(f"  Threshold: {calibration.get('current_threshold', 0):.4f}")
            doc.add_paragraph(f"  Accuracy: {calibration.get('accuracy', 0):.2%}")
            doc.add_paragraph(f"  Stable: {'Yes' if calibration.get('is_stable') else 'No'}")

    def _add_provenance_content(self, doc: Any, content: Dict[str, Any]) -> None:
        """Add provenance section content."""
        p = doc.add_paragraph()
        p.add_run("Stage: ").bold = True
        p.add_run(content.get("stage", "unknown"))

        p = doc.add_paragraph()
        p.add_run("Model: ").bold = True
        p.add_run(content.get("model", "unknown"))

        if content.get("created_at"):
            p = doc.add_paragraph()
            p.add_run("Created: ").bold = True
            p.add_run(content["created_at"])

        if content.get("pipeline_version"):
            p = doc.add_paragraph()
            p.add_run("Pipeline Version: ").bold = True
            p.add_run(content["pipeline_version"])

        # Reconstruction log
        recon = content.get("reconstruction")
        if recon:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Reconstruction Log: ").bold = True
            doc.add_paragraph(f"  Entries: {recon.get('entry_count', 0)}")
            doc.add_paragraph(f"  Current Phase: {recon.get('current_phase', 'N/A')}")

    def _add_generic_content(self, doc: Any, content: Any) -> None:
        """Add generic content (fallback for unhandled section types)."""
        if isinstance(content, dict):
            for key, value in content.items():
                if isinstance(value, (str, int, float, bool)):
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
        elif isinstance(content, str):
            doc.add_paragraph(content)

    def export_to_bytes(
        self,
        data: Any,
        options: ExportOptions,
    ) -> bytes:
        """Export data to DOCX bytes.

        Args:
            data: Data to export (AlignmentLayer or legacy types)
            options: Export options

        Returns:
            DOCX content as bytes
        """
        doc = self._create_document()

        # Check if AlignmentLayer (Stage E primary input)
        if self._is_alignment_layer(data):
            # Add title from AlignmentLayer
            self._add_title(doc, f"AlignmentLayer Export: {data.id}")

            # Add AlignmentLayer content using DocumentBuilder
            self._add_alignment_layer_content(doc, data, options)
        else:
            # Legacy handling for other data types
            title = "Pipeline Export"
            if hasattr(data, "image_id"):
                title = f"Export: {data.image_id}"
            elif hasattr(data, "export_id"):
                title = f"Export: {data.export_id}"

            self._add_title(doc, title)

            # Add content sections
            self._add_content_section(doc, data, options)

            # Add metadata
            self._add_metadata_section(doc, data, options)

            # Add provenance
            self._add_provenance_section(doc, data, options)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def export(
        self,
        data: Any,
        options: ExportOptions,
        image_id: str,
    ) -> ExportSpec:
        """Export data to DOCX file.

        Args:
            data: Data to export (AlignmentLayer or legacy types)
            options: Export options
            image_id: Source image identifier (or AlignmentLayer.id)

        Returns:
            ExportSpec with export metadata
        """
        start_time = time.time()

        # Generate identifiers
        export_id = self._generate_export_id(image_id)
        filename = self._generate_filename(image_id, options)
        filepath = self.config.output_dir / filename

        # Check existing file
        if filepath.exists() and not self.config.overwrite_existing:
            filepath = filepath.with_stem(f"{filepath.stem}_{export_id[:6]}")

        try:
            # Create document and export
            doc = self._create_document()

            # Check if AlignmentLayer (Stage E primary input)
            if self._is_alignment_layer(data):
                # Add title from AlignmentLayer
                self._add_title(doc, f"AlignmentLayer Export: {data.id}")

                # Add AlignmentLayer content using DocumentBuilder
                element_count = self._add_alignment_layer_content(doc, data, options)
            else:
                # Legacy handling
                self._add_title(doc, f"Export: {image_id}")

                # Add content and count elements
                element_count = self._add_content_section(doc, data, options)

                # Add metadata
                self._add_metadata_section(doc, data, options)

                # Add provenance
                self._add_provenance_section(doc, data, options)

            # Save to bytes for checksum
            buffer = BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()

            checksum = self._calculate_checksum(docx_bytes)
            file_size = self._write_file(docx_bytes, filepath)

            processing_time = (time.time() - start_time) * 1000

            self._stats["exports_completed"] += 1

            logger.info(
                f"DOCX export completed: {filepath}, "
                f"size={file_size} bytes, time={processing_time:.1f}ms"
            )

            return self._create_export_spec(
                export_id=export_id,
                image_id=image_id,
                filepath=filepath,
                file_size=file_size,
                checksum=checksum,
                element_count=element_count,
                processing_time_ms=processing_time,
            )

        except Exception as e:
            self._stats["exports_failed"] += 1
            logger.error(f"DOCX export failed: {e}")
            raise ExporterError(
                f"DOCX export failed: {e}",
                exporter_type="docx",
                element_id=image_id,
            )

    def validate_data(self, data: Any) -> bool:
        """Validate data before export.

        Args:
            data: Data to validate

        Returns:
            True if valid for DOCX export
        """
        if data is None:
            return False

        # AlignmentLayer (Stage E primary input)
        if self._is_alignment_layer(data):
            # Validate AlignmentLayer has required fields
            return bool(data.id and data.vision_spec_id)

        # Check for at least one exportable attribute (legacy)
        exportable_attrs = ["outputs", "nodes", "equations", "content", "alignments"]
        if any(hasattr(data, attr) for attr in exportable_attrs):
            return True

        # Accept dicts and lists
        if isinstance(data, (dict, list)):
            return True

        # Accept Pydantic models
        if isinstance(data, BaseModel):
            return True

        return False


# =============================================================================
# Export
# =============================================================================

__all__ = [
    "DOCXExporter",
    "DOCXExporterConfig",
]
