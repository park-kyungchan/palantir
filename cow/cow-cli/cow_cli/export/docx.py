"""
COW CLI - DOCX Exporter

Export separated documents to Microsoft Word format.
"""
from pathlib import Path
from typing import Optional

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from cow_cli.export.base import BaseExporter, ExportFormat, ExportOptions, ExportResult
from cow_cli.semantic.schemas import (
    SeparatedDocument,
    ContentElement,
    LayoutElement,
    ElementType,
)


class DocxExporter(BaseExporter):
    """
    DOCX format exporter.

    Creates a Microsoft Word document using python-docx.
    """

    format = ExportFormat.DOCX

    def __init__(
        self,
        options: Optional[ExportOptions] = None,
        template_path: Optional[Path] = None,
        include_toc: bool = False,
        math_as_image: bool = False,
    ):
        """
        Initialize DOCX exporter.

        Args:
            options: Export options
            template_path: Path to DOCX template
            include_toc: Include table of contents
            math_as_image: Render math as images instead of text
        """
        super().__init__(options)
        self.template_path = template_path
        self.include_toc = include_toc
        self.math_as_image = math_as_image

    def export(self, document: SeparatedDocument, output_path: Path) -> ExportResult:
        """
        Export document to DOCX.

        Args:
            document: The separated document
            output_path: Destination file path

        Returns:
            ExportResult with status
        """
        if not DOCX_AVAILABLE:
            return ExportResult(
                success=False,
                error="python-docx package not installed. Install with: pip install python-docx",
            )

        warnings = []
        output_path = self.ensure_extension(output_path)

        # Validate
        validation_errors = self.validate_document(document)
        if validation_errors:
            warnings.extend(validation_errors)

        try:
            # Create or load document
            if self.template_path and self.template_path.exists():
                doc = DocxDocument(str(self.template_path))
            else:
                doc = DocxDocument()

            # Build document content
            self._build_document(doc, document)

            # Save
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            return ExportResult(
                success=True,
                output_path=output_path,
                format=self.format,
                bytes_written=output_path.stat().st_size,
                warnings=warnings,
            )

        except Exception as e:
            return ExportResult(
                success=False,
                error=str(e),
                warnings=warnings,
            )

    def _build_document(self, doc: "DocxDocument", document: SeparatedDocument) -> None:
        """Build the DOCX document content."""
        # Title
        title = Path(document.image_path).stem.replace("_", " ").title()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Source info
        source_para = doc.add_paragraph()
        source_para.add_run("Source: ").bold = True
        source_para.add_run(document.image_path)
        source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # TOC placeholder
        if self.include_toc:
            self._add_toc_placeholder(doc)

        # Content sections
        current_section = None

        for content_elem in document.content.elements:
            layout_elem = document.layout.get_element_by_id(
                content_elem.layout_ref or content_elem.id
            )

            elem_type = layout_elem.type if layout_elem else ElementType.TEXT

            # Section headers
            if elem_type != current_section:
                section_name = self._format_type_name(elem_type)
                doc.add_heading(section_name, level=1)
                current_section = elem_type

            # Render element
            self._render_element(doc, content_elem, layout_elem)

        # Quality metrics
        if self.options.include_quality_metrics:
            self._add_quality_section(doc, document)

    def _format_type_name(self, elem_type) -> str:
        """Format element type for display."""
        if isinstance(elem_type, str):
            return elem_type.replace("_", " ").title()
        return str(elem_type).replace("_", " ").title()

    def _render_element(
        self,
        doc: "DocxDocument",
        content: ContentElement,
        layout: Optional[LayoutElement],
    ) -> None:
        """Render a single element to DOCX."""
        elem_type = layout.type if layout else ElementType.TEXT

        # Title
        if elem_type == ElementType.TITLE:
            self._render_title(doc, content)
            return

        # Section headers
        if elem_type == ElementType.SECTION_HEADER:
            self._render_section(doc, content)
            return

        # Math
        if elem_type == ElementType.MATH:
            self._render_math(doc, content)
            return

        # Table
        if elem_type == ElementType.TABLE:
            self._render_table(doc, content)
            return

        # Code
        if elem_type in (ElementType.CODE, ElementType.PSEUDOCODE):
            self._render_code(doc, content)
            return

        # Diagram/Chart
        if elem_type in (ElementType.DIAGRAM, ElementType.CHART):
            self._render_figure_placeholder(doc, content, layout)
            return

        # Default: text
        self._render_text(doc, content)

    def _render_title(self, doc: "DocxDocument", content: ContentElement) -> None:
        """Render title."""
        text = content.text or ""
        heading = doc.add_heading(text, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_section(self, doc: "DocxDocument", content: ContentElement) -> None:
        """Render section header."""
        text = content.text or ""
        doc.add_heading(text, level=2)

    def _render_math(self, doc: "DocxDocument", content: ContentElement) -> None:
        """Render math element."""
        latex = content.latex or content.latex_styled or ""

        if self.math_as_image:
            # Placeholder for math image
            para = doc.add_paragraph()
            para.add_run(f"[Math: {latex[:50]}...]" if len(latex) > 50 else f"[Math: {latex}]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Add as formatted text with equation style
            para = doc.add_paragraph()

            # Try to use OMML if available
            try:
                self._add_omml_math(para, latex)
            except Exception:
                # Fallback to plain text
                run = para.add_run(latex)
                run.font.name = "Cambria Math"
                run.font.size = Pt(11)

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_omml_math(self, paragraph, latex: str) -> None:
        """Add Office Math ML to paragraph."""
        # Create oMath element
        omath = OxmlElement("m:oMath")

        # Create a run with the LaTeX text
        omath_r = OxmlElement("m:r")
        omath_t = OxmlElement("m:t")
        omath_t.text = latex

        omath_r.append(omath_t)
        omath.append(omath_r)

        paragraph._p.append(omath)

    def _render_table(self, doc: "DocxDocument", content: ContentElement) -> None:
        """Render table element."""
        # Try TSV data first
        for data in content.data:
            if data.type.value == "tsv":
                self._tsv_to_docx_table(doc, data.value)
                return

        # Fallback to text
        text = content.text or ""
        para = doc.add_paragraph(text)
        para.style = "Quote"

    def _tsv_to_docx_table(self, doc: "DocxDocument", tsv: str) -> None:
        """Convert TSV to DOCX table."""
        lines = tsv.strip().split("\n")
        if not lines:
            return

        # Parse rows
        rows = [line.split("\t") for line in lines]
        num_rows = len(rows)
        num_cols = max(len(row) for row in rows)

        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        # Fill cells
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.text = cell_text

                    # Bold header row
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        doc.add_paragraph()  # Spacing

    def _render_code(self, doc: "DocxDocument", content: ContentElement) -> None:
        """Render code element."""
        code = content.text or ""

        para = doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

        # Add border/shading
        para_format = para.paragraph_format
        para_format.left_indent = Cm(0.5)
        para_format.right_indent = Cm(0.5)

    def _render_figure_placeholder(
        self,
        doc: "DocxDocument",
        content: ContentElement,
        layout: Optional[LayoutElement],
    ) -> None:
        """Render figure/diagram placeholder."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Placeholder text
        run = para.add_run("[Figure/Diagram]")
        run.italic = True

        # Caption if available
        if content.text:
            caption_para = doc.add_paragraph()
            caption_para.add_run(f"Caption: {content.text[:100]}")
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = "Caption"

        doc.add_paragraph()

    def _render_text(self, doc: "DocxDocument", content: ContentElement) -> None:
        """Render text element."""
        text = content.text or content.text_display or ""

        if not text and content.latex:
            # Inline math
            text = f"${content.latex}$"

        if text:
            doc.add_paragraph(text)

    def _add_toc_placeholder(self, doc: "DocxDocument") -> None:
        """Add table of contents placeholder."""
        doc.add_heading("Table of Contents", level=1)
        para = doc.add_paragraph()
        para.add_run("[TOC placeholder - Update field after opening in Word]")
        para.italic = True
        doc.add_page_break()

    def _add_quality_section(
        self,
        doc: "DocxDocument",
        document: SeparatedDocument,
    ) -> None:
        """Add quality metrics section."""
        doc.add_heading("Quality Metrics", level=1)

        summary = document.content.quality_summary

        # Create table
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        metrics = [
            ("Total Elements", str(summary.total_elements)),
            ("High Confidence", str(summary.high_confidence_count)),
            ("Needs Review", str(summary.needs_review_count)),
            ("Average Confidence", f"{summary.average_confidence:.1%}" if summary.average_confidence else "N/A"),
        ]

        for i, (label, value) in enumerate(metrics):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value


__all__ = ["DocxExporter", "DOCX_AVAILABLE"]
